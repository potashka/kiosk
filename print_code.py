import os
import logging
from docx import Document

# Настройка логирования
logging.basicConfig(
    level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def add_file_content_to_document(
        doc, directory, extensions, exclude_dirs=None, exclude_files=None
):
    if exclude_dirs is None:
        exclude_dirs = []
    if exclude_files is None:
        exclude_files = []

    for root, dirs, files in os.walk(directory):
        # Исключаем заданные директории из обхода
        dirs[:] = [d for d in dirs if d not in exclude_dirs]
        for file in files:
            file_path = os.path.join(root, file)
            # Проверяем, что файл соответствует заданным расширениям и не находится в списке исключений
            if file.endswith(tuple(extensions)) and os.path.basename(file) not in exclude_files:
                logger.info(f'Добавление файла: {file_path}')
                # Добавление названия директории и файла
                doc.add_heading(
                    'Файл: ' + file_path.replace(directory, ''), level=2
                )
                with open(file_path, 'r', encoding='utf-8') as f:
                    # Добавление содержимого файла
                    content = f.read()
                    paragraphs = content.split('\n')
                    for paragraph in paragraphs:
                        doc.add_paragraph(paragraph)
                    doc.add_page_break()  # Добавляем разрыв страницы после каждого файла

def create_word_document(
        directory, extensions,
        filename="IOTProjectCode.docx", exclude_dirs=None, exclude_files=None
):
    logger.info('Создание документа Word...')
    doc = Document()
    doc.add_heading('Код проекта IOT', 0)

    add_file_content_to_document(doc, directory, extensions, exclude_dirs, exclude_files)

    # Сохранение документа
    save_path = f"C:/dev/kiosk/{filename}"
    doc.save(save_path)
    logger.info(f"Документ '{filename}' успешно сохранён в {save_path}.")

# Пример использования
if __name__ == "__main__":
    directory = "C:\\dev\\kiosk\\"
    extensions = ['.py', '.html']  # Расширения файлов для включения
    exclude_dirs = ['venv', 'migrations', 'static']  # Список каталогов для исключения
    exclude_files = ['all_models.py', 'print_code.py', 'admin.py', 'main_old.py', '.gitignore']  # Список файлов для исключения
    create_word_document(
        directory, extensions, "IOTProjectCode.docx", exclude_dirs, exclude_files
    )
